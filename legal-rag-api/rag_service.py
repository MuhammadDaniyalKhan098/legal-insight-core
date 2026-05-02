"""
RAG Service Module - Core business logic
File: rag_service.py
"""

import os
import json
from typing import TypedDict, Literal, List, Dict, Any, Annotated
import base64
import io
from docx import Document

from langchain_core.messages import BaseMessage, HumanMessage, AIMessage
from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver
from langgraph.types import Send

from neo4j import GraphDatabase
import google.generativeai as genai
# from llama_index.embeddings.google_genai import GoogleGenAIEmbedding
from llama_index.core import Settings
from dotenv import load_dotenv
import re
import hashlib  # STRIDE -- Information Disclosure --> Hash sensitive query data for logging
import time  # STRIDE -- Denial of Service --> Track query timestamps for rate limiting

from llama_index.embeddings.jinaai import JinaEmbedding

def merge_retrieval_results(existing: List[Dict], new: List[Dict]) -> List[Dict]:
    """
    Reducer that combines results from parallel branches, 
    BUT allows clearing the history if a special signal is sent.
    """
    # Check if the new update is the special "CLEAR" signal
    if isinstance(new, list) and len(new) == 1 and new[0] == "__CLEAR__":
        return []
    
    # Otherwise, standard append behavior for parallel retrieval
    return existing + new

class LegalAgentState(TypedDict, total=False):
    messages: List[BaseMessage]
    query: str
    mode: str
    is_legal_query: bool
    is_comparison: bool
    is_prediction: bool  # <--- NEW: Added for Outcome Prediction Use Case
    is_procedure: bool
    comparison_terms: List[str] | None
    target_language: str
    sub_queries: List[str]
    retrieved_results: Annotated[List[Dict], merge_retrieval_results]
    context: str
    route: Literal["router", "retrieve", "generate", "casual_response", "end"]
    final_answer: str
    sub_query: str
    sub_query_idx: int
    db_target: int  # Added to track which DB to query (1 or 2)
    sub_query_embeddings: Dict[str, List[float]]  # <--- ADD THIS
    embedding_vector: List[float]                 # <--- ADD THIS The specific vector (Child)
    is_drafting: bool # New
    draft_template: Dict[str, Any] # New: Stores the retrieved template node
    user_provided_info: Dict[str, Any] # New: Info extracted from user prompt
    reasoning: Dict[str, Any]  # <--- New addition for drafting
    query_hash: str  # STRIDE -- Repudiation --> Track query hash for audit logging


class LegalRAGService:
    """Service class for Legal RAG functionality"""
    
    def __init__(self, env_path: str = ".env"):
        """Initialize the RAG service"""
        # Only load .env locally (Render already has env vars in system)
        if env_path and os.path.exists(env_path):
            load_dotenv(env_path)
        
        self.google_api_key = os.getenv("GOOGLE_API_KEY")
        
        # --- DB 1 Config (Laws/Statutes) ---
        self.db1_config = {
            "url": os.getenv("NEO4J_URL"),
            "user": os.getenv("NEO4J_USER"),
            "password": os.getenv("NEO4J_PASSWORD"),
            "database": os.getenv("NEO4J_DATABASE", "neo4j"),
            "label": "Statutory Law DB"
        }

        # --- DB 2 Config (Case Studies/Hearings) ---
        self.db2_config = {
            "url": os.getenv("NEO4J_URL_2"),
            "user": os.getenv("NEO4J_USER_2"),
            "password": os.getenv("NEO4J_PASSWORD_2"),
            "database": os.getenv("NEO4J_DATABASE_2", "neo4j"),
            "label": "Case Law & Hearings DB"
        }

        self.db3_config = {
            "url": os.getenv("NEO4J_URL_3"),
            "user": os.getenv("NEO4J_USER_3"),
            "password": os.getenv("NEO4J_PASSWORD_3"),
            "database": os.getenv("NEO4J_DATABASE_3", "neo4j"),
            "label": "Procedural & Drafting DB"
        }
        
        # --- DB 4 Config (Legal Templates) ---
        self.db4_config = {
            "url": os.getenv("NEO4J_URL_4"),
            "user": os.getenv("NEO4J_USER_4"),
            "password": os.getenv("NEO4J_PASSWORD_4"),
            "database": os.getenv("NEO4J_DATABASE_4", "neo4j"),
            "label": "Drafting Templates DB"
        }

        self.embed_model = None
        self.model = None
        self.legal_agent = None
        self._initialized = False
        
    def initialize(self):
        """Initialize models and graph"""
        if self._initialized:
            return
        
        # Validate environment variables
        if not self.google_api_key:
            raise ValueError("❌ GOOGLE_API_KEY not found in environment")
        if not self.db1_config["url"] or not self.db2_config["url"]:
            raise ValueError("❌ NEO4J URLs (1 or 2) not found in environment")
        
        # Initialize embedding model
        try:
            jina_key = os.getenv("JINA_API_KEY")
            if not jina_key:
                raise ValueError("❌ JINA_API_KEY not found in environment")
            self.embed_model = JinaEmbedding(
                api_key=jina_key,
                model="jina-embeddings-v3",
                task="retrieval.query" # Optimized for search queries
            )

            Settings.embed_model = self.embed_model
            print("✅ Embedding model initialized")
        except Exception as e:
            raise Exception(f"❌ Embedding model initialization failed: {e}")
        
        # Initialize Gemini LLM
        try:
            genai.configure(api_key=self.google_api_key)
            self.model = genai.GenerativeModel('models/gemini-2.5-flash')
            print("✅ Gemini LLM initialized")
        except Exception as e:
            raise Exception(f"❌ LLM initialization failed: {e}")
        
        # Initialize Neo4j schema for BOTH databases
        print("⚙️ Initializing schemas for both databases...")
        self._initialize_neo4j_schema(self.db1_config)
        self._initialize_neo4j_schema(self.db2_config)
        self._initialize_neo4j_schema(self.db3_config)
        
        # Build the graph
        self.legal_agent = self._build_legal_rag_graph()
        
        self._initialized = True
        print("✅ Legal RAG Service fully initialized")
    
    def is_ready(self) -> bool:
        """Check if service is ready"""
        return self._initialized and self.legal_agent is not None
    
    def _initialize_neo4j_schema(self, db_config: Dict):
        """Initialize Neo4j database schema for a specific config"""
        cypher_schema = [
            "CREATE CONSTRAINT sectionKey IF NOT EXISTS FOR (c:Section) REQUIRE (c.key) IS UNIQUE;",
            "CREATE CONSTRAINT chunkKey IF NOT EXISTS FOR (c:Chunk) REQUIRE (c.key) IS UNIQUE;",
            "CREATE CONSTRAINT documentKey IF NOT EXISTS FOR (c:Document) REQUIRE (c.url_hash) IS UNIQUE;",
            "CREATE VECTOR INDEX `chunkVectorIndex` IF NOT EXISTS FOR (e:Embedding) "
            "ON (e.value) OPTIONS { indexConfig: {`vector.dimensions`: 1024, `vector.similarity_function`: 'cosine'}};"
        ]
        
        driver = GraphDatabase.driver(
            db_config["url"],
            database=db_config["database"],
            auth=(db_config["user"], db_config["password"])
        )
        
        try:
            with driver.session() as session:
                for cypher in cypher_schema:
                    session.run(cypher)
            print(f"✅ Neo4j schema initialized for {db_config['label']}")
        except Exception as e:
            print(f"⚠️ Neo4j schema warning for {db_config['label']}: {e}")
        finally:
            driver.close()
    
    def _extract_relevant_nodes(self, query_embedding: List[float], db_config: Dict, limit: int = 5) -> List[Dict]:
        """Extract top relevant nodes from a specific Neo4j DB"""
        driver = GraphDatabase.driver(
            db_config["url"],
            database=db_config["database"],
            auth=(db_config["user"], db_config["password"])
        )
        
        try:
            with driver.session() as session:
                cypher_chunk_nodes = """
                    MATCH (c:Chunk)
                    WITH c, gds.similarity.cosine(c.embedding, $query_embedding) AS score
                    ORDER BY score DESC
                    LIMIT $limit
                    RETURN elementId(c) AS chunk_id, score
                """
                chunk_results = list(session.run(
                    cypher_chunk_nodes,
                    query_embedding=query_embedding,
                    limit=limit
                ))
                
                score_map = {record["chunk_id"]: record["score"] for record in chunk_results}
                chunk_ids = list(score_map.keys())
                
                if not chunk_ids:
                    return []
                
                cypher_relations = """
                    MATCH (c:Chunk)-[:UNDER_SECTION]->(s:Section)<-[:HAS_SECTION]-(d:Document)<-[:HAS_PAGE]-(p:PDF)
                    WHERE elementId(c) IN $chunk_ids
                    RETURN elementId(c) AS chunk_id,
                           c.text AS chunk_text,
                           s.text AS section_text,
                           p.name AS pdf_name
                """
                result_relations = session.run(cypher_relations, chunk_ids=chunk_ids)
                
                results = []
                for record in result_relations:
                    r_dict = dict(record)
                    r_dict["source_db"] = db_config["label"]
                
                    c_id = r_dict["chunk_id"]
                    r_dict["score"] = score_map.get(c_id, 0.0)
                    
                    results.append(r_dict)

                results.sort(key=lambda x: x["score"], reverse=True)
                
                return results
        
        except Exception as e:
            print(f"❌ Error retrieving nodes from {db_config['label']}: {e}")
            return []
        finally:
            driver.close()
    
    def _build_context(self, results: List[Dict], is_comparison: bool = False) -> str:
        """Build context from retrieval results"""
        if is_comparison:
            return self._build_context_for_comparison(results)
        return self._build_enhanced_context(results)
    
    def _build_context_for_comparison(self, all_results: List[Dict]) -> str:
        """Build structured context for comparison queries"""
        pdf_contexts = {}
        
        for r in all_results:
            if "[Omitted.]" in r['chunk_text'] or "[Omitted.]" in r['section_text']:
                continue
            
            # Key context by PDF name AND source DB to keep types distinct
            pdf_name = f"{r['pdf_name']} ({r.get('source_db', 'Unknown')})"
            
            if pdf_name not in pdf_contexts:
                pdf_contexts[pdf_name] = {'sections': set(), 'chunks': []}
            
            section_text = r['section_text']
            if section_text not in pdf_contexts[pdf_name]['sections']:
                pdf_contexts[pdf_name]['sections'].add(section_text)
            
            pdf_contexts[pdf_name]['chunks'].append({
                'chunk': r['chunk_text'],
                'section': section_text
            })
        
        context_parts = []
        for pdf_name, data in pdf_contexts.items():
            context_parts.append(f"=== Document: {pdf_name} ===")
            for chunk_info in data['chunks']:
                context_parts.append(f"Section: {chunk_info['section']}")
                context_parts.append(f"Content: {chunk_info['chunk']}")
                context_parts.append("")
        
        return "\n".join(context_parts)
    
    def _build_enhanced_context(self, results: List[Dict]) -> str:
        """Build enhanced context with sub-query awareness"""
        subquery_groups = {}
        for r in results:
            sq_idx = r.get('sub_query_idx', 0)
            subquery_groups.setdefault(sq_idx, []).append(r)
        
        context_parts = []
        
        for sq_idx in sorted(subquery_groups.keys()):
            results_group = subquery_groups[sq_idx]
            if not results_group:
                continue
            
            sub_query = results_group[0].get('sub_query', 'General')
            context_parts.append(f"\n=== Context for: {sub_query} ===\n")
            
            seen_sections = set()
            for r in results_group:
                if "[Omitted.]" in r['chunk_text'] or "[Omitted.]" in r['section_text']:
                    continue
                
                section_text = r['section_text']
                source_label = r.get('source_db', 'Unknown DB')
                
                if section_text not in seen_sections:
                    context_parts.append(f"Section: {section_text}")
                    seen_sections.add(section_text)
                
                # Explicitly mentioning source to prevent hallucination
                context_parts.append(f"Content: {r['chunk_text']}")
                context_parts.append(f"[Source: {r['pdf_name']} | Type: {source_label}]")
            
            context_parts.append("")
        
        return "\n".join(context_parts)
    
    def _build_legal_rag_graph(self):
        """Build the LangGraph with parallel retrieval"""
        graph = StateGraph(LegalAgentState)
        
        # Add nodes
        graph.add_node("router", self._router_node)
        graph.add_node("casual_response", self._casual_response_node)
        graph.add_node("analyze_and_decompose", self._analyze_and_decompose_node)
        graph.add_node("drafting", self._drafting_node) # NEW NODE
        graph.add_node("retrieve_single_subquery", self._retrieve_single_subquery_node)
        graph.add_node("merge_context", self._merge_and_build_context_node)
        graph.add_node("generator", self._generation_node)
        
        # Set entry point
        graph.set_entry_point("router")

        # graph.add_edge("router", "analyze_and_decompose")

        # 3. NEW: Conditional Routing from the entry point
        graph.add_conditional_edges(
            "router",
            lambda x: x["route"],
            {
                "drafting": "drafting",
                "analyze_and_decompose": "analyze_and_decompose"
            }
        )

        # 4. Standard Edges
        graph.add_conditional_edges(
            "analyze_and_decompose",
            self._route_to_parallel_retrieval,
            ["retrieve_single_subquery", "casual_response"] 
        )

        graph.add_edge("retrieve_single_subquery", "merge_context")
        graph.add_edge("merge_context", "generator")
        graph.add_edge("drafting", END) # Drafting completes here
        graph.add_edge("casual_response", END)
        graph.add_edge("generator", END)
        
        return graph.compile(checkpointer=MemorySaver())

    # --- Template Retrieval Logic ---
    def _retrieve_template(self, query_embedding: List[float]) -> Dict[str, Any]:
        driver = GraphDatabase.driver(
            self.db4_config["url"], 
            auth=(self.db4_config["user"], self.db4_config["password"])
        )
        try:
            with driver.session() as session:
                # Searching the LegalTemplate node from your screenshot
                cypher = """
                    MATCH (t:LegalTemplate)
                    WITH t, gds.similarity.cosine(t.embedding, $query_embedding) AS score
                    ORDER BY score DESC LIMIT 1
                    RETURN t.content as content, t.name as name, t.required_fields as fields, 
                           t.file_base64 as file_base64, t.filename as filename, score
                """
                result = session.run(cypher, query_embedding=query_embedding).single()
                return dict(result) if result else None
        finally:
            driver.close()

    def _sanitize_input(self, text: str, max_length: int = 10000) -> str:
        """
        STRIDE -- Tampering --> Sanitize user input to prevent injection attacks
        STRIDE -- Denial of Service --> Limit input length to prevent resource exhaustion
        """
        if not text:
            return ""
        
        # Limit length to prevent DoS
        sanitized = text[:max_length]
        
        # Remove potentially dangerous characters for template injection
        # This is a basic sanitization - adjust based on specific needs
        dangerous_patterns = [
            r'<script[^>]*>.*?</script>',  # Remove script tags
            r'javascript:',  # Remove javascript: protocol
            r'on\w+\s*=',  # Remove inline event handlers
        ]
        
        for pattern in dangerous_patterns:
            sanitized = re.sub(pattern, '', sanitized, flags=re.IGNORECASE | re.DOTALL)
        
        return sanitized.strip()

    def _validate_field_mappings(self, mappings: Dict[str, Any], template_fields: List[str]) -> Dict[str, Any]:
        """
        STRIDE -- Elevation of Privilege --> Validate that extracted fields match expected template fields
        STRIDE -- Tampering --> Prevent injection of unauthorized fields
        """
        validated = {}
        
        for key, value in mappings.items():
            # Only accept fields that are in the template's required fields
            if key in template_fields:
                # Sanitize the value
                if isinstance(value, str):
                    validated[key] = self._sanitize_input(value, max_length=500)
                else:
                    validated[key] = value
        
        return validated

    def _fill_docx_template(self, base64_content: str, data: Dict[str, Any]) -> str:
        """
        Dynamically fill a DOCX template with user data in memory.
        STRIDE -- Tampering --> Validate data before insertion to prevent malicious content injection
        """
        try:
            # 1. Decode base64 to bytes
            file_bytes = base64.b64decode(base64_content)
            doc_stream = io.BytesIO(file_bytes)
            doc = Document(doc_stream)

            # 2. Helper to replace text in a paragraph
            def replace_text_in_paragraph(paragraph, replacements):
                for key, value in replacements.items():
                    if key in paragraph.text:
                        # STRIDE -- Tampering --> Sanitize value before insertion
                        safe_value = self._sanitize_input(str(value), max_length=500)
                        
                        # Use Regex to handle {{ Key }} and {{Key}} variations
                        pattern = re.compile(r'\{\{\s*' + re.escape(key) + r'\s*\}\}', re.IGNORECASE)
                        if pattern.search(paragraph.text):
                            paragraph.text = pattern.sub(safe_value, paragraph.text)

            # 3. Iterate over all paragraphs
            for para in doc.paragraphs:
                replace_text_in_paragraph(para, data)

            # 4. Iterate over all tables (forms often use tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_text_in_paragraph(para, data)

            # 5. Save to new buffer
            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)

            # 6. Encode back to base64
            new_base64 = base64.b64encode(output_stream.read()).decode('utf-8')
            return new_base64

        except Exception as e:
            print(f"❌ Error generating DOCX: {e}")
            return base64_content # Fallback to original if filling fails

    def _get_all_template_names(self) -> List[str]:
        """Fetch a list of all available template names from DB 4"""
        driver = GraphDatabase.driver(
            self.db4_config["url"], 
            auth=(self.db4_config["user"], self.db4_config["password"])
        )
        try:
            with driver.session() as session:
                cypher = "MATCH (t:LegalTemplate) RETURN t.name as name ORDER BY t.name ASC"
                result = session.run(cypher)
                return [record["name"] for record in result]
        except Exception as e:
            print(f"❌ Error fetching template names: {e}")
            return []
        finally:
            driver.close()


    def _drafting_node(self, state: LegalAgentState) -> LegalAgentState:
        query = state["query"]
        template = state.get("draft_template")
        
        # STRIDE -- Tampering --> Sanitize query input before processing
        sanitized_query = self._sanitize_input(query, max_length=5000)
        
        # 1. Fetch template if not in state
        if not template:
            print("📝 No template in state, performing semantic search...")
            emb = self.embed_model.get_text_embedding(sanitized_query)
            template = self._retrieve_template(emb)
            if not template or template.get('score', 0) < 0.35:
                # return {**state, "final_answer": "Template not found.", "route": "end"}
                print(f"⚠️ Template rejected. Score: {template.get('score', 0)}")
                # Fetch available templates
                available_templates = self._get_all_template_names()

                # return {
                #     **state, 
                #     "final_answer": "I couldn't find a matching template for that request. Please be more specific or try a different document type.", 
                #     "route": "end"
                # }
                if available_templates:
                    # Format as a Markdown bulleted list
                    templates_list = "\n".join([f"- {name}" for name in available_templates])
                    ans = (
                        "I couldn't find a high-confidence match for that request. "
                        "Here are the templates currently available in my system:\n\n"
                        f"{templates_list}\n\n"
                        "**Please reply with one of these names.**"
                    )
                else:
                    ans = "I couldn't find a matching template, and the database appears to be empty."

                return {
                    **state, 
                    "final_answer": ans, 
                    "route": "end"
                }

            print(f"✅ Template matched: {template['name']} (Score: {template.get('score')})")

        mapping_prompt = f"""
        You are an expert legal document processor.
        
        Target Template Keys: {template['fields']}
        User Input: \"\"\"{sanitized_query}\"\"\"
        
        Task:
        1. Extract all personal and case details from the User Input.
        2. Map ONE user value to MULTIPLE target keys if they represent the same data.
        - Example: If the user provides a CNIC, map it to "LostCNICNumber", "LostCNICNumber,IfKnown", and "CNICNumber(ifAvailable)".
        - Example: If the user provides a City, map it to both "City" and "Place".
        3. For "SelectApplicable", determine if the user mentioned a 'Daily Diary' or 'FIR' and return that string.
        
        Return ONLY a JSON object: {{"key_from_target_fields": "value"}}
        """
            
        try:
            mapping_res = self.model.generate_content(mapping_prompt, generation_config=genai.types.GenerationConfig(temperature=0.0))
            res_text = mapping_res.text.strip()
            
            # Fallback for manual JSON extraction if the model includes markdown formatting
            if "```json" in res_text:
                res_text = res_text.split("```json")[1].split("```")[0].strip()
            elif "```" in res_text:
                res_text = res_text.split("```")[1].split("```")[0].strip()
                
            new_data = json.loads(res_text)
        except Exception as e:
            print(f"⚠️ Extraction fallback triggered: {e}")
            # Final fallback: Standard call without strict JSON config
            mapping_res = self.model.generate_content(mapping_prompt)
            new_data = {}

        print(f"📝 Successfully Extracted {len(new_data)} fields: {new_data}")

        # STRIDE -- Elevation of Privilege --> Validate extracted fields against template
        # STRIDE -- Tampering --> Sanitize extracted data
        validated_data = self._validate_field_mappings(new_data, template['fields'])
        print(f"✅ Validated {len(validated_data)} fields after security checks")

        # Merge with previous turn info
        current_info = state.get("user_provided_info", {})
        current_info.update(validated_data)

        # 3. Robust Substitution using Regex
        filled_content = template['content']
        for field, value in current_info.items():
            # This regex matches {{Field}}, {{ Field }}, {{  Field  }}, etc.
            pattern = re.compile(r'\{\{\s*' + re.escape(field) + r'\s*\}\}', re.IGNORECASE)
            # STRIDE -- Tampering --> Sanitize value before substitution
            safe_value = self._sanitize_input(str(value), max_length=500)
            filled_content = pattern.sub(safe_value, filled_content)

        # 4. Identify genuinely missing fields from the updated text
        # This finds any remaining {{...}} tags
        remaining = re.findall(r'\{\{\s*(.*?)\s*\}\}', filled_content)

        # --- NEW: GENERATE THE ACTUAL FILLED FILE ---
        # Only generate the file if we have data to fill
        final_file_b64 = template['file_base64']
        if current_info:
            print("💾 Generating filled DOCX file...")
            # STRIDE -- Tampering --> _fill_docx_template already sanitizes data
            final_file_b64 = self._fill_docx_template(template['file_base64'], current_info)
        # --------------------------------------------
        
        if remaining:
            missing_list = list(set(remaining))
            final_answer = f"### Draft Updated: {template['name']}\n\n{filled_content}\n\n**Missing Details:** {', '.join(missing_list)}"
            return {
                **state,
                "final_answer": final_answer,
                "draft_template": template, 
                "user_provided_info": current_info,
                "reasoning": {
                    "is_draft": True,
                    "filename": template['filename'],
                    "file_data": final_file_b64
                },
                "route": "end"
            }
        else:
            final_answer = f"### Finalized {template['name']}\n\n{filled_content}"
            return {
                **state,
                "final_answer": final_answer,
                "draft_template": None,       # <--- CLEAR THIS
                "user_provided_info": {},     # <--- CLEAR THIS
                "reasoning": {
                    "is_draft": True,
                    "filename": template['filename'],
                    "file_data": final_file_b64
                },
                "route": "end"
            }


    # --- Router Modification ---
    def _router_node(self, state: LegalAgentState) -> LegalAgentState:
        """Route initial query based on 'draft:' prefix"""
        query = next((m.content for m in reversed(state["messages"])
                      if isinstance(m, HumanMessage)), "").strip()
        
        # STRIDE -- Repudiation --> Generate query hash for audit trail
        query_hash = hashlib.sha256(query.encode()).hexdigest()[:16]
        mode = state.get("mode", "general")

        # Check for the explicit prefix
        if mode == "drafting":
            print("🏗️ DRAFTING PATH DETECTED")
            # clean_query = query[6:].strip() # Remove 'draft:'
            clean_query = query
            # STRIDE -- Tampering --> Sanitize query
            clean_query = self._sanitize_input(clean_query, max_length=5000)
            return {
                **state,
                "query": clean_query,
                "is_drafting": True,
                "route": "drafting",
                "reasoning": {},
                "query_hash": query_hash  # STRIDE -- Repudiation --> Store hash for logging
            }
        
        # Standard flow if prefix is missing
        return {
            **state, 
            "query": query, 
            "retrieved_results": [],
            "context": "",
            "is_drafting": False,
            "route": "analyze_and_decompose",
            "reasoning": {},
            "query_hash": query_hash  # STRIDE -- Repudiation --> Store hash for logging
        }
    
    def _casual_response_node(self, state: LegalAgentState) -> LegalAgentState:
        """
        Handle casual conversation, history references, or direct analysis tasks.
        Uses an internal LLM call to classify intent instead of hardcoded keywords.
        """
        print("\n" + "="*80)
        print("💬 CASUAL / DIRECT RESPONSE NODE")
        print("="*80 + "\n")

        query = state["query"]
        # STRIDE -- Tampering --> Sanitize query
        sanitized_query = self._sanitize_input(query, max_length=5000)
        
        messages = state.get("messages", [])
        
        # Build conversation context string
        conversation_context = ""
        if len(messages) > 1:
            recent_messages = messages[-11:-1] if len(messages) > 11 else messages[:-1]
            for msg in recent_messages:
                role = "User" if isinstance(msg, HumanMessage) else "Assistant"
                # STRIDE -- Information Disclosure --> Truncate context to prevent excessive data exposure
                content_preview = msg.content[:500] if len(msg.content) > 500 else msg.content
                conversation_context += f"{role}: {content_preview}\n\n"

        # ------------------------------------------------------------------
        # 1. INTENT CLASSIFICATION (LLM Call)
        # ------------------------------------------------------------------
        print("🧠 Classifying user intent via LLM...")
        
        classification_prompt = f"""
You are a routing assistant. Analyze the User Query below and classify it into exactly one of these three categories:

1. "DIRECT_TASK": The user has pasted a substantial text/paragraph IN THE CURRENT QUERY and wants you to analyze THAT SPECIFIC TEXT.
   - CRITICAL: Do NOT choose this if the user says "summarize this" but does not actually paste the text.
   - Examples: "Summarize this: [text...]", "What does this mean: [text...]"

2. "HISTORY_REF": The user is referring to or asking follow-up questions about text/content that was provided in PREVIOUS messages in this conversation. This includes when they refer to "the text I provided" or ask follow-up questions about previously discussed content WITHOUT providing the text again.
   - CRITICAL: Choose this if the user uses demonstrative pronouns like "this", "that", "it", or "the text" BUT does NOT paste new text in the query.
   - Examples: "Can you summarize this?", "What did you mean by that?", "Explain it again.", "Based on the text I provided, what was X?", "In that paragraph, what did Y mean?"
3. "CASUAL": Greetings, compliments, general questions, or small talk. (e.g., "Hi", "Thanks", "Who are you?").

CRITICAL GUIDELINES:
- If the user provides the actual text to analyze IN THIS QUERY → DIRECT_TASK
- If the user refers to text from previous messages WITHOUT providing it again → HISTORY_REF
- Pay close attention to phrases like "the text I provided", "that paragraph", "based on what you said", etc.

User Query: {sanitized_query}

Recent Conversation Context (for reference):
{conversation_context if conversation_context else "No recent context"}

Respond ONLY with the category name: DIRECT_TASK, HISTORY_REF, or CASUAL.
"""
        try:
            # Use a low temperature for strict classification
            cls_response = self.model.generate_content(
                classification_prompt, 
                generation_config=genai.types.GenerationConfig(temperature=0.0)
            )
            intent = cls_response.text.strip().upper()
            
            # Fallback cleanup in case the LLM is chatty
            if "DIRECT_TASK" in intent: intent = "DIRECT_TASK"
            elif "HISTORY_REF" in intent: intent = "HISTORY_REF"
            else: intent = "CASUAL"
            
        except Exception as e:
            print(f"⚠️ Classification failed: {e}. Defaulting to CASUAL.")
            intent = "CASUAL"

        print(f"👉 Classified Intent: {intent}")

        # ------------------------------------------------------------------
        # 2. GENERATE RESPONSE BASED ON INTENT
        # ------------------------------------------------------------------
        target_lang = state.get("target_language", "English")
        if intent == "DIRECT_TASK":
            # Scenario: User provided text in the prompt to analyze
            prompt = f"""You are an Expert Legal Simplifier and Translator.
The user has provided a snippet of legal text within their prompt.

Conversation History:
{conversation_context}

User Query/Text: {sanitized_query}
Target Language: {target_lang}

Your Task:
1. **Analyze:** Identify complex legal jargon and concepts in the provided text.
2. **Simplify & Translate:** - Rewrite the text in **{target_lang}**.
   - Replace jargon with simple, equivalent terms.
   - Restructure sentences for maximum clarity.
3. **Explain:** Provide a brief, bullet-pointed explanation of what the clause means in practice (e.g., "This means this rule overrides X").

Format your response as:
**Simplified Version ({target_lang}):**
[The simplified text]

**Practical Explanation:**
* [Point 1]
* [Point 2]
Response:"""

        elif intent == "HISTORY_REF":
            # Scenario: User is asking about previous context
            prompt = f"""You are a helpful legal assistant.
            
Full Conversation History:
{conversation_context}

User Request: {sanitized_query}

Instructions:
1. **Contextual Accuracy:** Answer strictly based on the conversation history provided above.
2. **Summarization Standards:** If the user asks to summarize a previous answer:
   - Do NOT give a one-sentence abstract.
   - Provide a **comprehensive summary** that retains all **key steps, legal sections, and actionable advice**.
   - Use **bullet points** to organize the summary clearly.
3. **Follow-ups:** If the user asks a specific question about a previous point (e.g., "What did step 3 mean?"), explain it in detail using the context.
4. Do not make up new laws; refer to what was already discussed.

Response:"""

        else:
            # Scenario: Casual / General (Default)
            recent_conv = f"Recent Conversation:\n{conversation_context}" if conversation_context else ""
            prompt = f"""You are a friendly legal assistant chatbot specializing in Pakistani law.

{recent_conv}

Current message: {sanitized_query}

Respond naturally and helpfully:
- For greetings: Greet warmly.
- For thanks: Acknowledge graciously.
- Keep responses concise (2-3 sentences).

Response:"""
        
        # Generate the final answer
        response = self.model.generate_content(prompt)
        answer = response.text.strip()

        print(f"💬 Response: {answer}\n")
        
        return {
            **state,
            "final_answer": answer,
            "messages": state["messages"] + [AIMessage(content=answer)],
            "route": "end"
        }

    def _analyze_and_decompose_node(self, state: LegalAgentState) -> LegalAgentState:
        """Analyze query and decompose if needed"""
        print("="*80)
        print("🔎 ANALYZE & DECOMPOSE NODE")
        print("="*80)

        query = next((m.content for m in reversed(state["messages"])
                      if isinstance(m, HumanMessage)), "")
        
        # STRIDE -- Tampering --> Sanitize query
        sanitized_query = self._sanitize_input(query, max_length=5000)

        mode = state.get("mode", "general")

        print(f"\n📝 Explicit Mode selected: {mode.upper()}\n")

        # 1. Map Explicit Modes directly to Graph Flags
        is_prediction = (mode == "prediction")
        is_procedure = (mode == "procedure")
        
        print(f"\n📝 Analyzing query: {sanitized_query[:100]}...\n")
        
        messages = state.get("messages", [])
        conversation_context = ""
        if len(messages) > 1:
            recent_messages = messages[-6:-1] if len(messages) > 6 else messages[:-1]
            for msg in recent_messages:
                role = "User" if isinstance(msg, HumanMessage) else "Assistant"
                # STRIDE -- Information Disclosure --> Limit context size
                content_preview = msg.content[:100] if len(msg.content) > 100 else msg.content
                conversation_context += f"{role}: {content_preview}...\n"
        recent_conv = f"Recent conversation:\n{conversation_context}\n" if conversation_context else ""
        
        analysis_prompt = f"""
You are a legal assistant query analyzer for Pakistani law.

Analyze the following user query and respond in JSON format with:

1. "is_legal": true/false
2. "is_comparison": true/false
5. "needs_decomposition": true/false
6. "subqueries": [list of strings]
7. "comparison_terms": [list of strings or null]
8. "target_language": "Urdu" or "English" (Default to "English" unless Urdu is explicitly requested)

CRITICAL RULES:
1. Classify as CASUAL (is_legal: false) if:
   - It is a greeting, meta-question, or vague request.
   - **THE USER PROVIDES THE CONTEXT/TEXT IN THE PROMPT.** (e.g., "Summarize this passage: [text]", "What is the main idea of: [text]", or "Analyze this: [text]").
   - Even if the text contains legal jargon, if the user PROVIDED the text, do NOT search the database.
2. Classify as LEGAL (is_legal: true) ONLY if:
   - The user is asking a question that requires searching EXTERNAL legal knowledge (laws/cases) NOT present in the prompt.
   - Specific Pakistani laws, acts, procedures, or cases are mentioned *without* the text being provided.
4. Detect comparisons (comparison, differ, versus, distinguish).
5. Decompose complex queries into 3-6 focused subqueries for retrieval.

{recent_conv}
Current query: {sanitized_query}
Respond ONLY in valid JSON.
"""
        
        try:
            response = self.model.generate_content(analysis_prompt)
            response_text = response.text.strip()
            
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0].strip()
            
            analysis = json.loads(response_text)
            
            is_legal = analysis.get("is_legal", False)
            is_comparison = analysis.get("is_comparison", False)
            target_language = analysis.get("target_language", "English")
            subqueries = analysis.get("subqueries") or [sanitized_query]
            comparison_terms = analysis.get("comparison_terms", None)

            # If it's a prediction, ensure we treat it as a legal query to trigger retrieval
            if is_procedure or is_prediction:
                is_legal = True
            
            # --- NEW: BATCH EMBEDDING START ---
            sub_query_embeddings_map = {}
            if is_legal and subqueries:
                print(f"⚡ Batch embedding {len(subqueries)} sub-queries to avoid rate limits...")
                try:
                    # Jina/LlamaIndex supports batching. 
                    # We pass the list of strings, getting a list of vectors.
                    embeddings = self.embed_model.get_text_embedding_batch(subqueries)
                    
                    # Map subquery text -> embedding vector
                    for sq_text, emb in zip(subqueries, embeddings):
                        sub_query_embeddings_map[sq_text] = emb
                        
                except Exception as e:
                    print(f"⚠️ Batch embedding failed: {e}")
                    # Fallback or empty map (will fail gracefully in retrieval)
            # --- NEW: BATCH EMBEDDING END ---

            print(f"📊 Mode Enforced: {mode.upper()}")
            print(f"📊 Classification: {'LEGAL' if is_legal else 'CASUAL'}")
            print(f"   Comparison: {is_comparison}")
            print(f"   Prediction: {is_prediction}")
            print(f"   Procedure: {is_procedure}")
            print(f"   Subqueries: {subqueries}")

            route = "route_to_parallel_retrieval" if is_legal else "casual_response"
            
            return {
                **state,
                "query": sanitized_query,
                "is_legal_query": is_legal,
                "is_comparison": is_comparison,
                "is_prediction": is_prediction, # <--- NEW
                "is_procedure": is_procedure,
                "target_language": target_language,
                "comparison_terms": comparison_terms,
                "sub_queries": subqueries,
                "sub_query_embeddings": sub_query_embeddings_map, # <--- STORE HERE
                "route": route,
                "retrieved_results": ["__CLEAR__"],
                "context": ""
            }
        
        except Exception as e:
            print(f"⚠️ Analysis error: {e}")
            return {
                **state,
                "query": sanitized_query,
                "is_legal_query": False,
                "sub_queries": [sanitized_query],
                "route": "casual_response",
                "retrieved_results": ["__CLEAR__"],
                "context": ""
            }
    
    
    def _route_to_parallel_retrieval(self, state: LegalAgentState) -> List[Send]:
        if state.get("route") == "casual_response":
            return "casual_response"

        sub_queries = state["sub_queries"]
        is_procedure = state.get("is_procedure", False) # <--- Check flag

        # --- NEW: Retrieve the map from the current state ---
        embeddings_map = state.get("sub_query_embeddings", {})
        # ----------------------------------------------------

        # Default targets: DB1 (Statutes) + DB2 (Cases)
        targets = [1, 2]
        
        # If procedure, ALSO query DB3 (Drafting/Procedure)
        if is_procedure:
            print("🚀 GUIDED PROCEDURE DETECTED: Including DB 3")
            targets.append(3)

        tasks = []
        for idx, sub_query in enumerate(sub_queries, 1):
            # --- NEW: Get the specific vector for this sub-query ---
            # We treat the text as the key.
            current_vector = embeddings_map.get(sub_query)
            # -------------------------------------------------------

            for db_id in targets:
                tasks.append(Send("retrieve_single_subquery", {
                    "sub_query": sub_query,
                    "sub_query_idx": idx,
                    "db_target": db_id,
                    "embedding_vector": current_vector  # <--- PASS DATA HERE
                }))
            
        return tasks
    
    def _retrieve_single_subquery_node(self, state: LegalAgentState) -> Dict[str, Any]:
        """Retrieve for single sub-query (runs in parallel)"""
        sub_query = state.get("sub_query", "")
        sub_query_idx = state.get("sub_query_idx", 0)
        db_target = state.get("db_target", 1) # Default to 1 if missing

        # It was passed directly in the Send payload
        sub_query_embedding = state.get("embedding_vector")
        
        if not sub_query_embedding:
            print(f"⚠️ No embedding passed for: {sub_query[:20]}... (Skipping)")
            return {"retrieved_results": []}
        # --------------------------------

        # Select config based on target
        if db_target == 1:
            target_config = self.db1_config
        elif db_target == 2:
            target_config = self.db2_config
        else:
            target_config = self.db3_config
        

        results = self._extract_relevant_nodes(sub_query_embedding, db_config=target_config, limit=15)
        
        tagged_results = []
        for result in results:
            result['sub_query'] = sub_query
            result['sub_query_idx'] = sub_query_idx
            # Note: 'source_db' was already added in _extract_relevant_nodes
            tagged_results.append(result)
        
        return {"retrieved_results": tagged_results}
    
    def _merge_and_build_context_node(self, state: LegalAgentState) -> LegalAgentState:
        """Merge parallel results and build context"""
        print("\n" + "="*80)
        print("🔗 MERGE & CONTEXT BUILDING NODE")
        print("="*80 + "\n")

        all_results = state.get("retrieved_results", [])
        
        seen_chunks = set()
        unique_results = []
        
        for result in all_results:
            chunk_id = result['chunk_text']
            if chunk_id not in seen_chunks:
                unique_results.append(result)
                seen_chunks.add(chunk_id)

        print(f"✅ Merged {len(unique_results)} unique chunks from parallel retrievals\n")
        
        is_comparison = state.get("is_comparison", False)
        context = self._build_context(unique_results, is_comparison)

        print(f"✅ Context built ({len(context)} characters)\n")
        
        return {**state, "context": context, "route": "generate"}
    
    def _generation_node(self, state: LegalAgentState) -> LegalAgentState:
        """Generate final answer"""
        print("\n" + "="*80)
        print("🤖 GENERATION NODE")
        print("="*80 + "\n")

        query = state["query"]
        context = state["context"]
        is_comparison = state.get("is_comparison", False)
        is_prediction = state.get("is_prediction", False) # <--- NEW
        messages = state.get("messages", [])
        
        # STRIDE -- Information Disclosure --> Limit conversation history to prevent data leakage
        conversation_history = ""
        if len(messages) > 1:
            recent_messages = messages[-6:-1] if len(messages) > 6 else messages[:-1]
            for msg in recent_messages:
                role = "User" if isinstance(msg, HumanMessage) else "Assistant"
                # Truncate message content to prevent excessive context
                content_preview = msg.content[:500] if len(msg.content) > 500 else msg.content
                conversation_history += f"{role}: {content_preview}\n"
        
        recent_conv = f"Previous Conversation:\n{conversation_history}\n" if conversation_history else ""

        # --- 1. Comparison Logic ---
        if is_comparison:
            llm_input = f"""You are a legal assistant specializing in Pakistani law.

{recent_conv}
Context: {context}

Question: {query}

Instructions:
- Highlight key similarities and differences
- Reference relevant provisions
- Structure clearly per document
- Consider conversation context

Answer:"""

        # --- 2. Outcome Prediction Logic (NEW) ---
        elif is_prediction:
            llm_input = f"""You are a Legal Outcome Predictor for the Pakistani legal system.
The user has provided case facts and requested a prediction of the outcome.

{recent_conv}
Context (Retrieved Precedents & Statutes): {context}

User Case Facts: {query}

Instructions:
1. Analyze the User's Case Facts against the Retrieved Precedents.
2. Predict the likely outcome (e.g., Conviction, Acquittal, Dismissal, Fine, Bail Grant/Reject).
3. Assign a Probability Estimate (e.g., "High likelihood of conviction (approx. 70-80%)").
4. Provide REASONS citing specific past judgments or statutes from the context.
5. If the retrieved context is insufficient to make a solid prediction, explicitly state: "Insufficient past data available to predict with confidence," and suggest relevant statutes instead.
6. Suggest "Alternate Courses" if applicable (e.g., "If the accused pleads guilty...").

Structure:
- **Predicted Outcome:** [Outcome & %]
- **Key Precedents:** [Cite cases from context]
- **Reasoning:** [Legal analysis]
- **Exceptions/Notes:** [Caveats]

Answer:"""
        # --- 3. Procedure Guidance Logic ---
        elif state.get("is_procedure"):
            llm_input = f"""
You are a Guided Legal Procedure Assistant. 
The user needs specific steps, drafting advice, or procedural guidance.
Context includes Statutes, Case Law, AND Procedural/Drafting specific data.

Instructions:
1. Provide a step-by-step checklist for the requested procedure.
2. If drafting is required, list key clauses or requirements found in the context.
3. Reference specific Forms, Sections, or Departments mentioned.
4. Mention any relevant Case Law (precedents) that affects the procedure.
5. Structure the answer clearly (e.g., "Step 1", "Required Documents", "Drafting Tips").

{recent_conv}
Context: {context}
User Query: {query}
Answer:"""

        # --- 4. Standard Legal Q&A Logic ---
        else:
            llm_input = f"""You are a legal assistant specializing in Pakistani law.

{recent_conv}
Context: {context}

Question: {query}

Instructions:
- Provide accurate, clear legal information
- Reference specific provisions when relevant
- Consider the conversation context when formulating your answer
- If the user is following up on a previous question, maintain continuity

Answer:"""
        
        response = self.model.generate_content(llm_input)
        answer = response.text

        print("✅ Answer generated\n")
        print(answer)
        
        return {
            **state,
            "final_answer": answer,
            "messages": state["messages"] + [AIMessage(content=answer)],
            "route": "end"
        }
    
    # Public methods
    def query(self, query: str, thread_id: str = "default", verbose: bool = False, mode: str = "general") -> Dict[str, Any]:
        """Query the legal RAG agent"""
        if not self.is_ready():
            raise Exception("Service not initialized")
        
        # STRIDE -- Denial of Service --> Track query timestamp for rate limiting (to be implemented at API level)
        query_timestamp = time.time()
        
        config = {"configurable": {"thread_id": thread_id}}
        
        try:
            current_state = self.legal_agent.get_state(config)
            existing_messages = current_state.values.get("messages", []) if current_state and current_state.values else []
        except:
            existing_messages = []
        
        all_messages = existing_messages + [HumanMessage(content=query)]
        
        result = self.legal_agent.invoke(
            # {"messages": all_messages},
            {"messages": all_messages, "mode": mode},
            config=config
        )
        
        # STRIDE -- Repudiation --> Log query hash for audit trail
        query_hash = result.get("query_hash", "")
        if verbose and query_hash:
            print(f"🔐 Query Hash (for audit): {query_hash}")
        
        return result
    
    def get_conversation_history(self, thread_id: str) -> List[Dict[str, Any]]:
        """Get conversation history for a thread"""
        config = {"configurable": {"thread_id": thread_id}}
        
        try:
            current_state = self.legal_agent.get_state(config)
            if not current_state or not current_state.values:
                return []
            
            messages = current_state.values.get("messages", [])
            return [
                {
                    "role": "user" if isinstance(m, HumanMessage) else "assistant",
                    "content": m.content
                }
                for m in messages
            ]
        except:
            return []
    
    def clear_conversation(self, thread_id: str):
        """Clear conversation history for a thread"""
        # LangGraph MemorySaver doesn't have direct clear method
        # This is a placeholder - history will naturally expire
        pass


# Convenience function for backward compatibility
def query_legal_agent(query: str, thread_id: str = "default", verbose: bool = True) -> str:
    """Standalone query function"""
    service = LegalRAGService()
    service.initialize()
    result = service.query(query, thread_id, verbose)
    return result.get("final_answer", "")